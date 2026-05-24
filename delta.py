# ==========================================
# DELTA | PREMIUM LEGAL ARCHITECTURE SUITE
# Draftable-Style Visual Overhaul Edition
# ==========================================

import streamlit as st
import streamlit.components.v1 as components
import difflib
import io
import re
import hashlib
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================================
# BLOCK 1: STATE HYDRATION
# ==========================================
if 'initialized' not in st.session_state:
    st.session_state.initialized = True
    st.session_state.uploaded_files_data = {}    # Document paragraph string lists
    st.session_state.uploaded_files_hashes = {}  # SHA-256 cryptographic hashes
    st.session_state.file_order = []
    st.session_state.file_roles = {}
    st.session_state.processing_complete = False
    st.session_state.current_baseline = 0
    st.session_state.current_counter = 1
    st.session_state.comparison_mode = "Baseline vs Counter"


# ==========================================
# BLOCK 2: GLOBAL LUXURY CSS — DRAFTABLE STYLE
# Injects the light-gray desktop, white paper
# sheets, hidden Streamlit chrome, and all
# inline highlight token styles.
# ==========================================
def inject_luxury_system_css():
    st.markdown("""
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@400;500&family=IBM+Plex+Serif:wght@300;400;500&family=IBM+Plex+Mono:wght@300;400&display=swap');

            /* ── Desktop Background ─────────────────────── */
            .stApp {
                background-color: #EAECEF !important;
                font-family: 'IBM Plex Serif', serif !important;
                color: #1A1A1A !important;
            }

            /* ── Hide Streamlit Chrome ───────────────────── */
            #MainMenu, footer, header,
            [data-testid="stToolbar"],
            [data-testid="stDecoration"] {
                display: none !important;
            }
            .block-container {
                padding-top: 1.5rem !important;
                padding-bottom: 0 !important;
                max-width: 100% !important;
            }

            /* ── Sidebar ─────────────────────────────────── */
            [data-testid="stSidebar"] {
                background-color: #F4F5F7 !important;
                border-right: 1px solid #D8D8D8 !important;
            }

            /* ── Brand Header Typography ─────────────────── */
            .brand-title {
                font-family: 'Cinzel', serif !important;
                color: #1A1A1A !important;
                letter-spacing: 5px;
                font-weight: 400;
                font-size: 26px;
                margin-bottom: 0.1rem;
            }
            .brand-subtitle {
                font-family: 'IBM Plex Mono', monospace;
                color: #888888;
                font-size: 9px;
                letter-spacing: 2.5px;
                text-transform: uppercase;
                margin-bottom: 2rem;
            }

            /* ── Landing Centering ───────────────────────── */
            .landing-wrapper {
                display: flex;
                flex-direction: column;
                align-items: center;
                justify-content: center;
                text-align: center;
                width: 100%;
                margin: 0 auto;
            }

            /* ── File Uploader ───────────────────────────── */
            .stFileUploader {
                border: 1px dashed #B0B0B0 !important;
                background-color: #FFFFFF !important;
                border-radius: 2px !important;
                padding: 1.5rem !important;
            }

            /* ── Buttons ─────────────────────────────────── */
            .stButton > button {
                background-color: transparent !important;
                color: #1A1A1A !important;
                border: 1px solid #1A1A1A !important;
                border-radius: 0px !important;
                font-family: 'IBM Plex Mono', monospace;
                font-size: 10px !important;
                letter-spacing: 1.5px;
                text-transform: uppercase;
                padding: 0.55rem 1.8rem !important;
                transition: all 0.2s ease;
                width: 100%;
            }
            .stButton > button:hover {
                background-color: #1A1A1A !important;
                color: #FFFFFF !important;
            }

            /* ── White Paper Document Viewport ───────────── */
            /* Applied to the continuous doc columns */
            .paper-viewport {
                background-color: #FFFFFF;
                color: #1A1A1A;
                padding: 30px 40px;
                border: 1px solid #E0E0E0;
                box-shadow: 0 4px 12px rgba(0,0,0,0.05);
                min-height: 100px;
                font-family: 'IBM Plex Serif', serif;
                font-size: 13px;
                line-height: 1.85;
            }

            /* ── Paragraph row inside the paper ─────────── */
            .paper-para {
                padding: 10px 0;
                border-bottom: 1px solid #F0F0F0;
                word-wrap: break-word;
            }
            .paper-para:last-child {
                border-bottom: none;
            }

            /* ── Empty placeholder block (alignment shim) ── */
            .paper-placeholder {
                background-color: #FAFAFA;
                border: 1px dashed #E8E8E8;
                border-radius: 2px;
                color: transparent;
                font-size: 13px;
                line-height: 1.85;
                padding: 10px 0;
                user-select: none;
            }

            /* ── Change type trace flags ─────────────────── */
            .trace-flag {
                font-family: 'IBM Plex Mono', monospace;
                font-size: 9px;
                letter-spacing: 1px;
                text-transform: uppercase;
                margin-bottom: 4px;
                display: block;
                color: #888888;
            }
            .trace-flag-del  { color: #991B1B; }
            .trace-flag-ins  { color: #065F46; }
            .trace-flag-mod  { color: #92400E; }

            /* ── Inline Diff Token Highlights ────────────── */
            .add-token {
                background-color: #D1FAE5;
                color: #065F46;
                border-radius: 2px;
                padding: 1px 3px;
            }
            .del-token {
                background-color: #FEE2E2;
                color: #991B1B;
                text-decoration: line-through;
                border-radius: 2px;
                padding: 1px 3px;
            }

            /* ── Mini-Map Center Track ───────────────────── */
            .minimap-track {
                display: flex;
                flex-direction: column;
                gap: 3px;
                align-items: center;
                padding-top: 8px;
            }
            .minimap-del {
                width: 10px;
                height: 8px;
                background-color: #FCA5A5;
                border-radius: 1px;
            }
            .minimap-ins {
                width: 10px;
                height: 8px;
                background-color: #6EE7B7;
                border-radius: 1px;
            }
            .minimap-eq {
                width: 10px;
                height: 8px;
                background-color: #D1D5DB;
                border-radius: 1px;
            }

            /* ── Executive Change-List Cards (Col 4) ──────── */
            .change-card {
                background-color: #FFFFFF;
                border: 1px solid #E5E7EB;
                border-radius: 2px;
                padding: 10px 12px;
                margin-bottom: 8px;
                font-family: 'IBM Plex Serif', serif;
            }
            .change-card-header {
                font-family: 'IBM Plex Mono', monospace;
                font-size: 9px;
                letter-spacing: 1px;
                text-transform: uppercase;
                color: #374151;
                margin-bottom: 4px;
                font-weight: 600;
            }
            .change-card-body {
                font-size: 11px;
                color: #6B7280;
                line-height: 1.5;
            }
            .change-badge-del {
                display: inline-block;
                background-color: #FEE2E2;
                color: #991B1B;
                font-family: 'IBM Plex Mono', monospace;
                font-size: 9px;
                padding: 1px 5px;
                border-radius: 2px;
                margin-right: 4px;
                font-weight: 700;
            }
            .change-badge-ins {
                display: inline-block;
                background-color: #D1FAE5;
                color: #065F46;
                font-family: 'IBM Plex Mono', monospace;
                font-size: 9px;
                padding: 1px 5px;
                border-radius: 2px;
                margin-right: 4px;
                font-weight: 700;
            }
            .change-badge-mod {
                display: inline-block;
                background-color: #FEF3C7;
                color: #92400E;
                font-family: 'IBM Plex Mono', monospace;
                font-size: 9px;
                padding: 1px 5px;
                border-radius: 2px;
                margin-right: 4px;
                font-weight: 700;
            }

            /* ── Crypto Manifest Bar ─────────────────────── */
            .crypto-banner {
                font-family: 'IBM Plex Mono', monospace;
                font-size: 9px;
                background-color: #FFFFFF;
                border: 1px solid #E5E7EB;
                padding: 6px 14px;
                color: #9CA3AF;
                margin-bottom: 1rem;
                border-radius: 2px;
            }

            /* ── Column Header Labels ────────────────────── */
            .col-label {
                font-family: 'IBM Plex Mono', monospace;
                font-size: 9px;
                text-transform: uppercase;
                letter-spacing: 2px;
                color: #6B7280;
                font-weight: 600;
                padding-bottom: 6px;
                border-bottom: 1px solid #D1D5DB;
                margin-bottom: 8px;
            }

            /* ── Select / input overrides for light theme ─── */
            div[data-baseweb="select"] {
                background-color: #FFFFFF !important;
                border-radius: 0px !important;
            }
            input {
                border-radius: 0px !important;
                background-color: #FFFFFF !important;
                color: #1A1A1A !important;
                border: 1px solid #D1D5DB !important;
                font-size: 12px !important;
                padding: 6px 10px !important;
            }

            /* ── Dividers ────────────────────────────────── */
            hr { border-color: #E5E7EB !important; }
        </style>
    """, unsafe_allow_html=True)


# ==========================================
# BLOCK 3: ENGINE PARSING & CRYPTO HASHING
# ==========================================
def parse_pdf(file_bytes):
    """Extract paragraph-level text blocks from a PDF using PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    paragraphs = []
    for page in doc:
        blocks = page.get_text("blocks")
        for b in blocks:
            text = b[4].strip()
            if text:
                paragraphs.append(" ".join(text.split()))
    return paragraphs


def parse_docx(file_bytes):
    """Extract non-empty paragraphs from a .docx document."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(" ".join(p.text.split()))
    return paragraphs


def load_staged_matrices(uploaded_files):
    """Parse and cache uploaded documents; compute SHA-256 hashes."""
    for file in uploaded_files:
        if file.name not in st.session_state.uploaded_files_data:
            bytes_data = file.read()

            # [ADVANCED FEATURE 3: SHA-256 Cryptographic Hash Verification]
            file_hash = hashlib.sha256(bytes_data).hexdigest()
            st.session_state.uploaded_files_hashes[file.name] = file_hash

            if file.name.endswith('.pdf'):
                parsed_text = parse_pdf(bytes_data)
            elif file.name.endswith('.docx'):
                parsed_text = parse_docx(bytes_data)
            else:
                continue

            st.session_state.uploaded_files_data[file.name] = parsed_text
            if file.name not in st.session_state.file_order:
                st.session_state.file_order.append(file.name)
                st.session_state.file_roles[file.name] = "v1: Baseline"


# ==========================================
# BLOCK 4: FUZZY ALIGNMENT ENGINE
# ==========================================
# [ADVANCED FEATURE 2: Position-agnostic token-ratio bipartite alignment]
def compute_fuzzy_alignment_matrix(left_paras, right_paras, threshold=0.45):
    """
    Aligns paragraphs between two documents using fuzzy token-sort matching.
    Returns a list of opcodes: (tag, i1, i2, j1, j2) analogous to difflib format.
    """
    matched_right_indices = set()
    alignment_opcodes = []

    for i, lp in enumerate(left_paras):
        best_ratio = 0.0
        best_j = None
        lp_tokens = sorted(lp.lower().split())  # Token sort neutralizes re-ordering

        for j, rp in enumerate(right_paras):
            if j in matched_right_indices:
                continue
            rp_tokens = sorted(rp.lower().split())
            ratio = difflib.SequenceMatcher(None, lp_tokens, rp_tokens).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_j = j

        if best_ratio >= threshold and best_j is not None:
            matched_right_indices.add(best_j)
            if best_ratio > 0.98 and lp == right_paras[best_j]:
                alignment_opcodes.append(('equal', i, i, best_j, best_j))
            else:
                alignment_opcodes.append(('replace', i, i, best_j, best_j))
        else:
            alignment_opcodes.append(('delete', i, i, None, None))

    # Collect unmatched right paragraphs as insertions
    for j in range(len(right_paras)):
        if j not in matched_right_indices:
            alignment_opcodes.append(('insert', None, None, j, j))

    # Sort to preserve the structural execution flow of the baseline document
    alignment_opcodes.sort(
        key=lambda x: (
            x[1] if x[1] is not None else float('inf'),
            x[3] if x[3] is not None else 0
        )
    )
    return alignment_opcodes


# ==========================================
# BLOCK 5: STRING TRANSFORMATION ANALYTICS
# ==========================================
def extract_clause_signature(text):
    """Extract a short clause identifier from the beginning of a paragraph."""
    match = re.match(r'^([A-Za-z0-9\.\s]+(?:\b[A-Z]{2,}\b|\bRent\b|\bTerm\b|\bDeposit\b|\bUse\b))', text)
    if match:
        return match.group(1).strip()
    words = text.split()
    return " ".join(words[:3]) if len(words) >= 3 else "Provision Block"


def generate_advisory_trace_text(text1, text2, signature):
    """Generate a human-readable summary of the word-level diff between two paragraphs."""
    matcher = difflib.SequenceMatcher(None, text1.split(), text2.split())
    traces = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace':
            o = " ".join(text1.split()[i1:i2])
            n = " ".join(text2.split()[j1:j2])
            if len(o) < 30 and len(n) < 30:
                traces.append(f"Changed '{o}' → '{n}'")
        elif tag == 'delete':
            d = " ".join(text1.split()[i1:i2])
            if len(d) < 30:
                traces.append(f"Omitted '{d}'")
        elif tag == 'insert':
            ins = " ".join(text2.split()[j1:j2])
            if len(ins) < 30:
                traces.append(f"Injected '{ins}'")
    return f"{signature} | " + ("; ".join(traces) if traces else "Terms modified.")


def compute_token_diff_html(text1, text2):
    """
    Computes word-level diff between two strings and returns two HTML strings
    with inline <span> highlights for additions and deletions.
    """
    matcher = difflib.SequenceMatcher(None, text1.split(), text2.split())
    out1, out2 = [], []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        w1 = " ".join(text1.split()[i1:i2])
        w2 = " ".join(text2.split()[j1:j2])
        if tag == 'equal':
            out1.append(w1)
            out2.append(w2)
        elif tag == 'replace':
            out1.append(f'<span class="del-token">{w1}</span>')
            out2.append(f'<span class="add-token">{w2}</span>')
        elif tag == 'delete':
            out1.append(f'<span class="del-token">{w1}</span>')
        elif tag == 'insert':
            out2.append(f'<span class="add-token">{w2}</span>')
    return " ".join(out1), " ".join(out2)


# ==========================================
# BLOCK 6: HIGH-FIDELITY COMPLIANCE EXPORTERS
# ==========================================
def set_run_background(run, color_hex):
    """Apply a Word background shading colour to a run element."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    rPr.append(shd)


def export_landscape_docx(left_paras, right_paras, title_left, title_right, alignment_opcodes, hash_l, hash_r):
    """Generate a landscape .docx review matrix with highlighted diffs."""
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    title = doc.add_paragraph()
    title.add_run("DELTA CONTRACT ADVISORY MATRIX").bold = True

    # Embed SHA-256 integrity hashes in document header
    meta = doc.add_paragraph()
    meta.add_run(
        f"INTEGRITY MANIFEST LOG\nBASE SHA-256: {hash_l}\nCNTR SHA-256: {hash_r}\n"
    ).font.size = Pt(8)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = f"BASELINE ({title_left})"
    hdr_cells[1].text = f"COUNTERPART ({title_right})"
    hdr_cells[2].text = "SMART DELTA EVALUATION"

    for tag, i1, _, j1, _ in alignment_opcodes:
        row = table.add_row()
        if tag == 'equal':
            row.cells[0].text = left_paras[i1]
            row.cells[1].text = right_paras[j1]
            row.cells[2].text = "No variance detected."

        elif tag == 'replace':
            p1 = row.cells[0].paragraphs[0]
            p2 = row.cells[1].paragraphs[0]
            m_words = difflib.SequenceMatcher(
                None, left_paras[i1].split(), right_paras[j1].split()
            )
            for w_tag, w_i1, w_i2, w_j1, w_j2 in m_words.get_opcodes():
                w1_str = " ".join(left_paras[i1].split()[w_i1:w_i2]) + " "
                w2_str = " ".join(right_paras[j1].split()[w_j1:w_j2]) + " "
                if w_tag == 'equal':
                    p1.add_run(w1_str)
                    p2.add_run(w2_str)
                else:
                    r1 = p1.add_run(w1_str)
                    set_run_background(r1, "fee2e2")   # Pastel Red
                    r2 = p2.add_run(w2_str)
                    set_run_background(r2, "dcfce7")   # Pastel Green

            act_val = st.session_state.get(f"act_mod_{i1}_{j1}", "Unassigned")
            note_val = st.session_state.get(f"note_mod_{i1}_{j1}", "")
            sig = extract_clause_signature(left_paras[i1])
            trace = generate_advisory_trace_text(left_paras[i1], right_paras[j1], sig)
            row.cells[2].text = f"Evaluation: {trace}\nAction: {act_val}\nComments: {note_val}"

        elif tag == 'delete':
            r1 = row.cells[0].paragraphs[0].add_run(left_paras[i1])
            set_run_background(r1, "fee2e2")
            row.cells[1].text = "[Clause Omitted]"
            act_val = st.session_state.get(f"act_del_{i1}", "Unassigned")
            note_val = st.session_state.get(f"note_del_{i1}", "")
            row.cells[2].text = f"Evaluation: Structural Omission\nAction: {act_val}\nComments: {note_val}"

        elif tag == 'insert':
            row.cells[0].text = "[Absent from Baseline Template]"
            r2 = row.cells[1].paragraphs[0].add_run(right_paras[j1])
            set_run_background(r2, "dcfce7")
            act_val = st.session_state.get(f"act_ins_{j1}", "Unassigned")
            note_val = st.session_state.get(f"note_ins_{j1}", "")
            row.cells[2].text = f"Evaluation: Structural Insertion\nAction: {act_val}\nComments: {note_val}"

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def export_landscape_pdf(left_paras, right_paras, title_left, title_right, alignment_opcodes, hash_l, hash_r):
    """Generate a landscape PDF review matrix using PyMuPDF."""
    doc = fitz.open()
    page_w, page_h = 792, 612
    page = doc.new_page(width=page_w, height=page_h)

    page.insert_text(fitz.Point(36, 30), "DELTA CONTRACT REVIEW MATRIX", fontsize=12, color=(0.1, 0.1, 0.1))
    page.insert_text(
        fitz.Point(36, 45),
        f"BASE SHA-256: {hash_l} | CNTR SHA-256: {hash_r}",
        fontsize=7, color=(0.55, 0.55, 0.55)
    )

    y = 75
    col_w = 230
    c1_x, c2_x, c3_x = 36, 280, 524

    for tag, i1, _, j1, _ in alignment_opcodes:
        if y > (page_h - 90):
            page = doc.new_page(width=page_w, height=page_h)
            y = 50

        if tag == 'equal':
            page.insert_textbox(fitz.Rect(c1_x, y, c1_x + col_w, y + 65), left_paras[i1], fontsize=8, color=(0.1, 0.1, 0.1))
            page.insert_textbox(fitz.Rect(c2_x, y, c2_x + col_w, y + 65), right_paras[j1], fontsize=8, color=(0.1, 0.1, 0.1))
            page.insert_textbox(fitz.Rect(c3_x, y, c3_x + col_w, y + 65), "No variance detected.", fontsize=8, color=(0.5, 0.5, 0.5))
            y += 70

        elif tag == 'replace':
            page.draw_rect(fitz.Rect(c1_x - 4, y, c1_x + col_w + 4, y + 70), color=None, fill=(0.99, 0.94, 0.94))
            page.draw_rect(fitz.Rect(c2_x - 4, y, c2_x + col_w + 4, y + 70), color=None, fill=(0.92, 0.99, 0.95))
            page.insert_textbox(fitz.Rect(c1_x, y + 5, c1_x + col_w, y + 65), left_paras[i1], fontsize=8, color=(0.4, 0.1, 0.1))
            page.insert_textbox(fitz.Rect(c2_x, y + 5, c2_x + col_w, y + 65), right_paras[j1], fontsize=8, color=(0.1, 0.3, 0.1))
            act_val = st.session_state.get(f"act_mod_{i1}_{j1}", "Unassigned")
            note_val = st.session_state.get(f"note_mod_{i1}_{j1}", "")
            sig = extract_clause_signature(left_paras[i1])
            trace = generate_advisory_trace_text(left_paras[i1], right_paras[j1], sig)
            page.insert_textbox(
                fitz.Rect(c3_x, y + 5, c3_x + col_w, y + 65),
                f"Eval: {trace}\nAction: {act_val}\nComments: {note_val}",
                fontsize=8, color=(0.1, 0.1, 0.1)
            )
            y += 75

        elif tag == 'delete':
            page.draw_rect(fitz.Rect(c1_x - 4, y, c1_x + col_w + 4, y + 55), color=None, fill=(0.99, 0.94, 0.94))
            page.insert_textbox(fitz.Rect(c1_x, y + 5, c1_x + col_w, y + 50), left_paras[i1], fontsize=8, color=(0.4, 0.1, 0.1))
            page.insert_text(fitz.Point(c2_x, y + 15), "[Clause Omitted]", fontsize=8, color=(0.5, 0.5, 0.5))
            act_val = st.session_state.get(f"act_del_{i1}", "Unassigned")
            note_val = st.session_state.get(f"note_del_{i1}", "")
            page.insert_textbox(
                fitz.Rect(c3_x, y + 5, c3_x + col_w, y + 50),
                f"Eval: Structural Omission\nAction: {act_val}\nComments: {note_val}",
                fontsize=8, color=(0.1, 0.1, 0.1)
            )
            y += 60

        elif tag == 'insert':
            page.draw_rect(fitz.Rect(c2_x - 4, y, c2_x + col_w + 4, y + 55), color=None, fill=(0.92, 0.99, 0.95))
            page.insert_text(fitz.Point(c1_x, y + 15), "[Absent from Baseline Template]", fontsize=8, color=(0.5, 0.5, 0.5))
            page.insert_textbox(fitz.Rect(c2_x, y + 5, c2_x + col_w, y + 50), right_paras[j1], fontsize=8, color=(0.1, 0.3, 0.1))
            act_val = st.session_state.get(f"act_ins_{j1}", "Unassigned")
            note_val = st.session_state.get(f"note_ins_{j1}", "")
            page.insert_textbox(
                fitz.Rect(c3_x, y + 5, c3_x + col_w, y + 50),
                f"Eval: Structural Insertion\nAction: {act_val}\nComments: {note_val}",
                fontsize=8, color=(0.1, 0.1, 0.1)
            )
            y += 60

    return doc.write()


# ==========================================
# BLOCK 7: UI PHASE 1 — CENTERED LANDING PORTAL
# ==========================================
def render_premium_landing_view():
    """
    Full-page centered landing screen for file ingestion.
    Light-theme: gray desktop, white uploader card, black typography.
    """
    st.markdown('<div class="landing-wrapper">', unsafe_allow_html=True)
    st.markdown('<div class="brand-title">DELTA</div>', unsafe_allow_html=True)
    st.markdown('<div class="brand-subtitle">Legal Infrastructure Pipeline</div>', unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "Ingest Core Binaries",
        type=['pdf', 'docx'],
        accept_multiple_files=True,
        label_visibility="collapsed"
    )
    st.markdown('</div>', unsafe_allow_html=True)

    if uploaded_files:
        load_staged_matrices(uploaded_files)
        st.markdown("<div style='max-width:500px; margin: 0 auto;'>", unsafe_allow_html=True)

        for idx, filename in enumerate(st.session_state.file_order):
            col1, col2 = st.columns([6, 4])
            with col1:
                st.markdown(
                    f'<p style="font-size:13px; color:#1A1A1A; padding-top:8px;">◼ {filename}</p>',
                    unsafe_allow_html=True
                )
            with col2:
                roles = ["Standard Template", "v1: Baseline", "v2: Counter", "v3: Counter", "v4: Counter"]
                st.session_state.file_roles[filename] = st.selectbox(
                    f"Classification_{filename}", roles,
                    index=roles.index(st.session_state.file_roles.get(filename, "v1: Baseline")),
                    label_visibility="collapsed",
                    key=f"stage_role_{filename}"
                )

        st.markdown("<br/>", unsafe_allow_html=True)
        if st.button("Compile Revision Tree"):
            st.session_state.processing_complete = True
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)


# ==========================================
# BLOCK 8: UI PHASE 2 — DRAFTABLE-STYLE REVIEW MATRIX
# ==========================================
def render_delta_contracts_view():
    """
    Four-column Draftable-style split view:
      Col 1 (4)   — Left white paper document viewport (Baseline)
      Col 2 (0.5) — Centre mini-map change tracker
      Col 3 (4)   — Right white paper document viewport (Counter)
      Col 4 (3.5) — Executive change list sidebar
    """

    # ── Header ──────────────────────────────────────────────────────────────
    st.markdown('<div class="brand-title" style="font-size:20px;">DELTA CONTRACTS</div>', unsafe_allow_html=True)
    st.markdown("<br/>", unsafe_allow_html=True)

    # ── Document & role resolution ───────────────────────────────────────────
    ordered_files = st.session_state.file_order
    roles = [st.session_state.file_roles[f] for f in ordered_files]

    base_file    = ordered_files[st.session_state.current_baseline]
    counter_file = ordered_files[st.session_state.current_counter]
    template_file = next(
        (f for f in ordered_files if "Standard Template" in st.session_state.file_roles[f]),
        None
    )

    # ── Document selector controls ───────────────────────────────────────────
    t_col1, t_col2 = st.columns(2)
    with t_col1:
        st.session_state.current_baseline = st.selectbox(
            "Baseline Selection Frame",
            range(len(ordered_files)),
            format_func=lambda x: f"BASE // {roles[x]}"
        )
    with t_col2:
        st.session_state.current_counter = st.selectbox(
            "Counterparty Selection Frame",
            range(len(ordered_files)),
            format_func=lambda x: f"CNTR // {roles[x]}",
            index=min(1, len(ordered_files) - 1)
        )

    # ── Sidebar: Standard Matrix reference track ─────────────────────────────
    with st.sidebar:
        st.markdown(
            '<div class="brand-title" style="font-size:15px; margin-top:1rem;">Vault</div>',
            unsafe_allow_html=True
        )
        if template_file:
            st.markdown(
                f'<p style="color:#888888; font-size:11px;">Matrix Reference: {template_file}</p>',
                unsafe_allow_html=True
            )
            modes = ["Baseline vs Counter", "Standard vs Baseline", "Standard vs Counter"]
            st.session_state.comparison_mode = st.radio(
                "Target Pipeline Strategy Routing",
                modes,
                index=modes.index(st.session_state.comparison_mode)
            )
            st.markdown("<hr/>", unsafe_allow_html=True)
            for p in st.session_state.uploaded_files_data[template_file]:
                st.markdown(
                    f'<div style="font-size:12px; color:#555555; margin-bottom:1rem; line-height:1.5;">{p}</div>',
                    unsafe_allow_html=True
                )
        else:
            st.markdown(
                '<p style="color:#888888; font-size:12px;">Standard Matrix signature absent in workspace staging array.</p>',
                unsafe_allow_html=True
            )

    st.markdown("<br/>", unsafe_allow_html=True)

    # ── Route left/right paragraph arrays based on comparison mode ───────────
    if st.session_state.comparison_mode == "Standard vs Baseline":
        left_paras  = st.session_state.uploaded_files_data[template_file] if template_file else st.session_state.uploaded_files_data[base_file]
        right_paras = st.session_state.uploaded_files_data[base_file]
        hash_l      = st.session_state.uploaded_files_hashes.get(template_file, "N/A")
        hash_r      = st.session_state.uploaded_files_hashes.get(base_file, "N/A")
        col1_title, col2_title = "STANDARD MATRIX", roles[st.session_state.current_baseline]

    elif st.session_state.comparison_mode == "Standard vs Counter":
        left_paras  = st.session_state.uploaded_files_data[template_file] if template_file else st.session_state.uploaded_files_data[base_file]
        right_paras = st.session_state.uploaded_files_data[counter_file]
        hash_l      = st.session_state.uploaded_files_hashes.get(template_file, "N/A")
        hash_r      = st.session_state.uploaded_files_hashes.get(counter_file, "N/A")
        col1_title, col2_title = "STANDARD MATRIX", roles[st.session_state.current_counter]

    else:  # Baseline vs Counter (default)
        left_paras  = st.session_state.uploaded_files_data[base_file]
        right_paras = st.session_state.uploaded_files_data[counter_file]
        hash_l      = st.session_state.uploaded_files_hashes.get(base_file, "N/A")
        hash_r      = st.session_state.uploaded_files_hashes.get(counter_file, "N/A")
        col1_title, col2_title = roles[st.session_state.current_baseline], roles[st.session_state.current_counter]

    # ── [ADVANCED FEATURE 2] Run Fuzzy Alignment Matrix ─────────────────────
    alignment_opcodes = compute_fuzzy_alignment_matrix(left_paras, right_paras)

    # ── [ADVANCED FEATURE 3] Cryptographic Validation Manifest Banner ────────
    st.markdown(f"""
        <div class="crypto-banner">
            VALIDATION MANIFEST LOG // LOCKED STATUS &nbsp;|&nbsp;
            [BASE_SIG]: {hash_l[:24]}…&nbsp;
            [CNTR_SIG]: {hash_r[:24]}…
        </div>
    """, unsafe_allow_html=True)

    # ── 4-Column Draftable Layout: [Left Doc | MiniMap | Right Doc | Changes] ─
    h_col1, h_col2, h_col3, h_col4 = st.columns([4, 0.5, 4, 3.5])

    with h_col1:
        st.markdown(f'<div class="col-label">{col1_title}</div>', unsafe_allow_html=True)
    with h_col2:
        st.markdown('<div class="col-label" style="text-align:center;">▲</div>', unsafe_allow_html=True)
    with h_col3:
        st.markdown(f'<div class="col-label">{col2_title}</div>', unsafe_allow_html=True)
    with h_col4:
        st.markdown('<div class="col-label">SMART DELTA EVALUATION</div>', unsafe_allow_html=True)

    # ── Build HTML buffers for the white paper viewports ─────────────────────
    # We accumulate all paragraph rows into HTML strings so we can render
    # a single continuous white-page block per document column, rather than
    # fragmented per-row Streamlit containers.

    left_html_rows   = []   # HTML snippets for each row in the left doc
    right_html_rows  = []   # HTML snippets for each row in the right doc
    minimap_html     = []   # Tiny bar blocks for the center mini-map
    change_cards     = []   # Dicts for the right-side change-list sidebar
    change_idx       = 0    # Sequential change counter for sidebar labels

    # ── Inline action/comment widgets live outside the HTML buffer ────────────
    # Because Streamlit's HTML blocks cannot contain interactive widgets, we
    # collect the unique_ids for 'replace', 'delete', 'insert' rows and render
    # the radio + text_input widgets separately inside column 4 after
    # the paper viewports are flushed.
    interactive_rows = []   # List of dicts with widget metadata for col 4

    # ── Loop over all alignment opcodes and build row HTML ───────────────────
    for tag, i1, _, j1, _ in alignment_opcodes:

        # ────────────────────────────────────────
        # EQUAL rows — unchanged text, plain para
        # ────────────────────────────────────────
        if tag == 'equal':
            para_text = left_paras[i1]
            left_html_rows.append(
                f'<div class="paper-para">{para_text}</div>'
            )
            right_html_rows.append(
                f'<div class="paper-para">{right_paras[j1]}</div>'
            )
            minimap_html.append('<div class="minimap-eq"></div>')

        # ────────────────────────────────────────
        # REPLACE rows — inline token diff highlights
        # ────────────────────────────────────────
        elif tag == 'replace':
            change_idx += 1
            unique_id   = f"mod_{i1}_{j1}"
            signature   = extract_clause_signature(left_paras[i1])
            trace_detail = generate_advisory_trace_text(left_paras[i1], right_paras[j1], signature)
            h1, h2      = compute_token_diff_html(left_paras[i1], right_paras[j1])

            left_html_rows.append(f"""
                <div class="paper-para">
                    <span class="trace-flag trace-flag-mod">▲ Modified</span>
                    {h1}
                </div>
            """)
            right_html_rows.append(f"""
                <div class="paper-para">
                    <span class="trace-flag trace-flag-mod">▲ Modified</span>
                    {h2}
                </div>
            """)
            minimap_html.append('<div class="minimap-del"></div>')
            minimap_html.append('<div class="minimap-ins"></div>')

            change_cards.append({
                'idx': change_idx,
                'badge': 'mod',
                'badge_label': '±',
                'header': f"{change_idx}. CHANGED // {signature}",
                'body': trace_detail,
                'unique_id': unique_id,
            })
            interactive_rows.append({
                'unique_id': unique_id,
                'type': 'replace',
                'i1': i1, 'j1': j1,
            })

        # ────────────────────────────────────────
        # DELETE rows — left has text, right is a placeholder shim
        # ────────────────────────────────────────
        elif tag == 'delete':
            change_idx += 1
            unique_id   = f"del_{i1}"
            signature   = extract_clause_signature(left_paras[i1])

            left_html_rows.append(f"""
                <div class="paper-para">
                    <span class="trace-flag trace-flag-del">◼ Omitted</span>
                    <span class="del-token">{left_paras[i1]}</span>
                </div>
            """)
            # Right side: visually aligned empty placeholder block
            right_html_rows.append(
                f'<div class="paper-placeholder">{left_paras[i1]}</div>'
            )
            minimap_html.append('<div class="minimap-del"></div>')

            change_cards.append({
                'idx': change_idx,
                'badge': 'del',
                'badge_label': '−',
                'header': f"{change_idx}. REMOVED // {signature}",
                'body': "Provision completely absent from this iteration track.",
                'unique_id': unique_id,
            })
            interactive_rows.append({
                'unique_id': unique_id,
                'type': 'delete',
                'i1': i1, 'j1': None,
            })

        # ────────────────────────────────────────
        # INSERT rows — right has text, left is a placeholder shim
        # ────────────────────────────────────────
        elif tag == 'insert':
            change_idx += 1
            unique_id   = f"ins_{j1}"
            signature   = extract_clause_signature(right_paras[j1])

            # Left side: visually aligned empty placeholder block
            left_html_rows.append(
                f'<div class="paper-placeholder">{right_paras[j1]}</div>'
            )
            right_html_rows.append(f"""
                <div class="paper-para">
                    <span class="trace-flag trace-flag-ins">◆ Insertion</span>
                    <span class="add-token">{right_paras[j1]}</span>
                </div>
            """)
            minimap_html.append('<div class="minimap-ins"></div>')

            change_cards.append({
                'idx': change_idx,
                'badge': 'ins',
                'badge_label': '+',
                'header': f"{change_idx}. ADDED // {signature}",
                'body': "New provision injected into this iteration track.",
                'unique_id': unique_id,
            })
            interactive_rows.append({
                'unique_id': unique_id,
                'type': 'insert',
                'i1': None, 'j1': j1,
            })

    # ── Render the four-column Draftable layout ──────────────────────────────
    doc_col1, map_col, doc_col3, sidebar_col = st.columns([4, 0.5, 4, 3.5])

    # Shared inline CSS injected into every components.html() iframe so styles
    # are self-contained and bypass Streamlit's markdown HTML sanitizer.
    IFRAME_STYLES = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Serif:wght@300;400;500&family=IBM+Plex+Mono:wght@300;400&display=swap');
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body { background: transparent; font-family: 'IBM Plex Serif', serif; font-size: 13px; color: #1A1A1A; line-height: 1.85; }
        .paper-viewport {
            background-color: #FFFFFF;
            padding: 30px 40px;
            border: 1px solid #E0E0E0;
            box-shadow: 0 4px 12px rgba(0,0,0,0.05);
            min-height: 200px;
        }
        .paper-para {
            padding: 10px 0;
            border-bottom: 1px solid #F0F0F0;
            word-wrap: break-word;
        }
        .paper-para:last-child { border-bottom: none; }
        .paper-placeholder {
            background-color: #FAFAFA;
            border: 1px dashed #E8E8E8;
            border-radius: 2px;
            color: transparent;
            padding: 10px 0;
            user-select: none;
        }
        .trace-flag {
            font-family: 'IBM Plex Mono', monospace;
            font-size: 9px;
            letter-spacing: 1px;
            text-transform: uppercase;
            margin-bottom: 4px;
            display: block;
            color: #888888;
        }
        .trace-flag-del { color: #991B1B; }
        .trace-flag-ins { color: #065F46; }
        .trace-flag-mod { color: #92400E; }
        .add-token {
            background-color: #D1FAE5;
            color: #065F46;
            border-radius: 2px;
            padding: 1px 3px;
        }
        .del-token {
            background-color: #FEE2E2;
            color: #991B1B;
            text-decoration: line-through;
            border-radius: 2px;
            padding: 1px 3px;
        }
        /* minimap styles */
        .minimap-track {
            display: flex;
            flex-direction: column;
            gap: 3px;
            align-items: center;
            padding-top: 8px;
        }
        .minimap-del { width: 10px; height: 8px; background-color: #FCA5A5; border-radius: 1px; }
        .minimap-ins { width: 10px; height: 8px; background-color: #6EE7B7; border-radius: 1px; }
        .minimap-eq  { width: 10px; height: 8px; background-color: #D1D5DB; border-radius: 1px; }
    </style>
    """

    # Estimate total paragraph count to size the iframe height dynamically
    total_rows = len(alignment_opcodes)
    estimated_height = max(400, total_rows * 90)

    # ── COLUMN 1: Left white paper viewport (iframe) ─────────────────────────
    with doc_col1:
        combined_left = "\n".join(left_html_rows)
        components.html(
            f"{IFRAME_STYLES}<div class='paper-viewport'>{combined_left}</div>",
            height=estimated_height,
            scrolling=True
        )

    # ── COLUMN 2: Centre mini-map visual track (iframe) ──────────────────────
    with map_col:
        combined_minimap = "\n".join(minimap_html)
        components.html(
            f"{IFRAME_STYLES}<div class='minimap-track'>{combined_minimap}</div>",
            height=estimated_height,
            scrolling=False
        )

    # ── COLUMN 3: Right white paper viewport (iframe) ────────────────────────
    with doc_col3:
        combined_right = "\n".join(right_html_rows)
        components.html(
            f"{IFRAME_STYLES}<div class='paper-viewport'>{combined_right}</div>",
            height=estimated_height,
            scrolling=True
        )

    # ── COLUMN 4: Executive change-list sidebar ──────────────────────────────
    # Renders one card per changed row with badge, header, trace detail,
    # then Streamlit-native action radio + comment input below each card.
    with sidebar_col:
        if not change_cards:
            st.markdown(
                '<div class="change-card"><div class="change-card-body" style="color:#9CA3AF;">'
                'No structural variance detected between selected documents.</div></div>',
                unsafe_allow_html=True
            )
        else:
            for card, widget_meta in zip(change_cards, interactive_rows):
                unique_id = card['unique_id']
                badge_class = f"change-badge-{card['badge']}"

                st.markdown(f"""
                    <div class="change-card">
                        <div class="change-card-header">
                            <span class="{badge_class}">{card['badge_label']}</span>
                            {card['header']}
                        </div>
                        <div class="change-card-body">{card['body']}</div>
                    </div>
                """, unsafe_allow_html=True)

                # Action protocol radio (Streamlit widget — must live outside HTML)
                st.radio(
                    "Action Protocol",
                    ["✓ Accept", "⚠ Escalate", "⇄ Counter", "✕ Reject"],
                    key=f"act_{unique_id}",
                    horizontal=True,
                    label_visibility="collapsed",
                    index=None,
                    help="Select a disposition for this change"
                )
                st.text_input(
                    "Comments",
                    key=f"note_{unique_id}",
                    placeholder="Add a comment…",
                    label_visibility="collapsed"
                )

    # ── Export footer ────────────────────────────────────────────────────────
    st.markdown("<br/><br/><hr/>", unsafe_allow_html=True)
    b_col1, b_col2 = st.columns(2)

    with b_col1:
        target_format = st.selectbox(
            "Export Options",
            ["Microsoft Word (.docx)", "Adobe Portable Document (.pdf)"],
            label_visibility="collapsed"
        )
        if "Word" in target_format:
            export_bytes = export_landscape_docx(
                left_paras, right_paras, col1_title, col2_title,
                alignment_opcodes, hash_l, hash_r
            )
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            extension = "docx"
        else:
            export_bytes = export_landscape_pdf(
                left_paras, right_paras, col1_title, col2_title,
                alignment_opcodes, hash_l, hash_r
            )
            mime_type = "application/pdf"
            extension = "pdf"

        st.download_button(
            label="📥 Export Review Matrix",
            data=export_bytes,
            file_name=f"DELTA_Matrix_{datetime.now().strftime('%Y%m%d')}.{extension}",
            mime=mime_type
        )

    with b_col2:
        st.write("<div style='height:1px;'></div>", unsafe_allow_html=True)
        if st.button("Terminate Session"):
            st.session_state.processing_complete = False
            st.rerun()


# ==========================================
# BLOCK 9: ENTRYPOINT ORCHESTRATION
# ==========================================
def main():
    st.set_page_config(
        page_title="DELTA CONTRACTS",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    inject_luxury_system_css()

    if not st.session_state.processing_complete:
        render_premium_landing_view()
    else:
        render_delta_contracts_view()


if __name__ == "__main__":
    main()
